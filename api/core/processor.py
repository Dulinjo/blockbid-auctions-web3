from __future__ import annotations

import io
import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import fitz
from docx import Document as DocxDocument
from odf import teletype
from odf.opendocument import load as odf_load

BASE_DIR = Path(__file__).resolve().parents[2]


def get_runtime_data_dir() -> Path:
    # Vercel serverless filesystem is read-only except /tmp.
    # Use /tmp in hosted runtime, repo data/ in local development.
    return Path("/tmp/lexvibe") if os.getenv("VERCEL") else BASE_DIR / "data"


class DocumentProcessingError(Exception):
    """Raised when parsing or validation of an uploaded file fails."""


@dataclass(slots=True)
class StoredDocument:
    filename: str
    normalized_text: str
    original_length: int


@dataclass(slots=True)
class DecisionMetadata:
    court: str
    decision_number: str


CYRILLIC_TO_LATIN_MAP = {
    "А": "A",
    "а": "a",
    "Б": "B",
    "б": "b",
    "В": "V",
    "в": "v",
    "Г": "G",
    "г": "g",
    "Д": "D",
    "д": "d",
    "Ђ": "Đ",
    "ђ": "đ",
    "Е": "E",
    "е": "e",
    "Ж": "Ž",
    "ж": "ž",
    "З": "Z",
    "з": "z",
    "И": "I",
    "и": "i",
    "Ј": "J",
    "ј": "j",
    "К": "K",
    "к": "k",
    "Л": "L",
    "л": "l",
    "Љ": "Lj",
    "љ": "lj",
    "М": "M",
    "м": "m",
    "Н": "N",
    "н": "n",
    "Њ": "Nj",
    "њ": "nj",
    "О": "O",
    "о": "o",
    "П": "P",
    "п": "p",
    "Р": "R",
    "р": "r",
    "С": "S",
    "с": "s",
    "Т": "T",
    "т": "t",
    "Ћ": "Ć",
    "ћ": "ć",
    "У": "U",
    "у": "u",
    "Ф": "F",
    "ф": "f",
    "Х": "H",
    "х": "h",
    "Ц": "C",
    "ц": "c",
    "Ч": "Č",
    "ч": "č",
    "Џ": "Dž",
    "џ": "dž",
    "Ш": "Š",
    "ш": "š",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".odt"}


def ensure_documents_dir() -> Path:
    runtime_data_dir = get_runtime_data_dir()
    documents_dir = runtime_data_dir / "documents"
    index_dir = runtime_data_dir / "index"
    documents_dir.mkdir(parents=True, exist_ok=True)
    index_dir.mkdir(parents=True, exist_ok=True)
    return documents_dir


def _sanitize_filename(filename: str) -> str:
    safe_name = Path(filename).name
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", safe_name)
    if not safe_name or "." not in safe_name:
        raise DocumentProcessingError("Neispravan naziv dokumenta.")
    return safe_name


def _slug_to_title(value: str) -> str:
    parts = [part for part in re.split(r"[-_]+", value) if part]
    if not parts:
        return "Nepoznati sud"
    stop_words = {"i", "u", "na", "od", "za", "o"}
    titled_parts: list[str] = []
    for index, part in enumerate(parts):
        cleaned = part.strip()
        if not cleaned:
            continue
        if index > 0 and cleaned in stop_words:
            titled_parts.append(cleaned)
        else:
            titled_parts.append(cleaned.capitalize())
    return " ".join(titled_parts) if titled_parts else "Nepoznati sud"


def extract_decision_metadata(filename: str) -> DecisionMetadata:
    stem = Path(filename).stem.lower()
    tokens = [token for token in re.split(r"[-_]+", stem) if token]

    for idx in range(0, len(tokens) - 2):
        code = tokens[idx]
        number = tokens[idx + 1]
        year = tokens[idx + 2]
        if not re.search(r"[a-z]", code):
            continue
        if not number.isdigit():
            continue
        if not re.fullmatch(r"\d{4}", year):
            continue

        decision_number = f"{code.upper()}-{number}/{year}"
        court_slug = "-".join(tokens[:idx]) if idx > 0 else "nepoznati-sud"
        return DecisionMetadata(
            court=_slug_to_title(court_slug),
            decision_number=decision_number,
        )

    return DecisionMetadata(court="Nepoznati sud", decision_number=Path(filename).stem)


def normalize_serbian_text(text: str) -> str:
    transliterated = "".join(CYRILLIC_TO_LATIN_MAP.get(char, char) for char in text)
    transliterated = transliterated.replace("ﬁ", "fi").replace("ﬂ", "fl")
    transliterated = unicodedata.normalize("NFKC", transliterated)
    normalized = re.sub(r"[ \t]+", " ", transliterated)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    return normalized.strip()


def _extract_pdf(file_bytes: bytes) -> str:
    parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)


def _extract_odt(file_bytes: bytes) -> str:
    doc = odf_load(io.BytesIO(file_bytes))
    return teletype.extractText(doc.text)


EXTRACTORS: dict[str, Callable[[bytes], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".odt": _extract_odt,
}


def parse_and_normalize_file(filename: str, file_handle: io.BufferedReader | io.BytesIO) -> StoredDocument:
    safe_filename = _sanitize_filename(filename)
    extension = Path(safe_filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentProcessingError(
            f"Nepodržan format dokumenta: {extension}. Dozvoljeni su PDF, DOCX i ODT."
        )

    raw_bytes = file_handle.read()
    if not raw_bytes:
        raise DocumentProcessingError("Dokument je prazan.")

    extractor = EXTRACTORS[extension]
    extracted_text = extractor(raw_bytes).strip()
    if not extracted_text:
        raise DocumentProcessingError("Dokument ne sadrži prepoznatljiv tekst.")

    normalized_text = normalize_serbian_text(extracted_text)
    if len(normalized_text) < 30:
        raise DocumentProcessingError("Dokument je prekratak za indeksiranje.")

    return StoredDocument(
        filename=safe_filename,
        normalized_text=normalized_text,
        original_length=len(extracted_text),
    )


def persist_upload(filename: str, raw_bytes: bytes) -> Path:
    documents_dir = ensure_documents_dir()
    safe_filename = _sanitize_filename(filename)
    destination = documents_dir / safe_filename
    destination.write_bytes(raw_bytes)
    return destination
